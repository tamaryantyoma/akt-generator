#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Генератор актов приема-передачи техники (Word)
Автор: Артём Тамарян
"""

import os
import json
from datetime import date
from docx import Document

# Tkinter используется только для прямого запуска app.py (не в веб-режиме)
# На PythonAnywhere tkinter недоступен — импорт безопасно пропускается
try:
    from tkinter import Tk, StringVar, Toplevel, messagebox, simpledialog
    from tkinter import ttk
    TKINTER_AVAILABLE = True
except ImportError:
    TKINTER_AVAILABLE = False

# ─── Пути ───────────────────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(BASE_DIR, "akt_handum_erevan_bilingual_no_borders_V2.docx")
MODELS_FILE = os.path.join(BASE_DIR, "laptop_models.json")
OUTPUT_DIR = os.path.join(BASE_DIR, "generated_acts")

# ─── Месяцы ─────────────────────────────────────────────────────────────────
ARMENIAN_MONTHS = [
    "հունվարի", "փետրվարի", "մարտի", "ապրիլի",
    "մայիսի", "հունիսի", "հուլիսի", "օգոստոսի",
    "սեպտեմբերի", "հոկտեմբերի", "նոյեմբերի", "դեկտեմբերի"
]

RUSSIAN_MONTHS = [
    "января", "февраля", "марта", "апреля",
    "мая", "июня", "июля", "августа",
    "сентября", "октября", "ноября", "декабря"
]


# ─── Загрузка / сохранение моделей ──────────────────────────────────────────
def load_models():
    """Загружает список моделей из JSON-файла."""
    if not os.path.exists(MODELS_FILE):
        return []
    with open(MODELS_FILE, "r", encoding="utf-8") as f:
        data = json.load(f)
    return data.get("models", [])


def save_models(models):
    """Сохраняет список моделей в JSON-файл."""
    with open(MODELS_FILE, "w", encoding="utf-8") as f:
        json.dump({"models": sorted(set(models))}, f, ensure_ascii=False, indent=2)


# ─── Генерация документа ────────────────────────────────────────────────────
def generate_act(name_arm, name_rus, gender, citizenship, model, serial, output_path):
    """
    Создаёт Word-документ на основе шаблона с подстановкой данных.
    
    Параметры:
        name_arm     — ФИО на армянском (Անուն Հայրանուն Ազգանուն)
        name_rus     — ФИО на русском (Фамилия Имя Отчество)
        gender       — "Мужской" или "Женский"
        citizenship  — "РФ" (Россия) или "РА" (Армения)
        model        — модель ноутбука
        serial       — серийный номер
        output_path  — куда сохранить
    """
    doc = Document(TEMPLATE_PATH)
    table = doc.tables[0]

    # Разбираем ФИО
    arm_parts = name_arm.strip().split()
    rus_parts = name_rus.strip().split()

    # Если частей меньше 3, дополняем пустыми строками
    while len(arm_parts) < 3:
        arm_parts.append("")
    while len(rus_parts) < 3:
        rus_parts.append("")

    arm_first, arm_patr, arm_last = arm_parts[0], arm_parts[1], arm_parts[2]
    rus_last, rus_first, rus_patr = rus_parts[0], rus_parts[1], rus_parts[2]

    today = date.today()
    day = today.day
    month_num = today.month
    year = today.year

    arm_month = ARMENIAN_MONTHS[month_num - 1]
    rus_month = RUSSIAN_MONTHS[month_num - 1]
    citizen_word = "гражданин" if gender == "Мужской" else "гражданка"

    # ── Cell [0,0] — дата (армянская часть) ──
    cell_00 = table.cell(0, 0)
    for run in cell_00.paragraphs[3].runs:
        if run.text == "01":
            run.text = str(day).zfill(2)
        elif run.text == "հունիսի":
            run.text = arm_month
        elif run.text == "2026":
            run.text = str(year)

    # ── Cell [0,1] — дата (русская часть) ──
    cell_01 = table.cell(0, 1)
    for run in cell_01.paragraphs[3].runs:
        if run.text == "01":
            run.text = str(day).zfill(2)
        elif run.text == "июня":
            run.text = rus_month
        elif run.text == " 2026":
            run.text = f" {year}"

    # ── Cell [1,0] — армянская часть (ФИО + модель + серийник) ──
    cell_10 = table.cell(1, 0)

    # Имя на армянском
    for run in cell_10.paragraphs[2].runs:
        if run.text == "Լեոնիդ ":
            run.text = f"{arm_first} "
        elif run.text == "Ստանիսլավի":
            run.text = arm_patr
        elif run.text == "Սիտրակով":
            run.text = arm_last

    # Гражданство (армянская сторона) — ՌԴ → ՀՀ для РА
    for run in cell_10.paragraphs[1].runs:
        if run.text == "ՌԴ" and citizenship == "РА":
            run.text = "ՀՀ"

    # Модель (армянская сторона) — paragraph[5], bold-ран после «Մոդել՝ »
    # Находим первый не-bold ран (это «Մոդել՝ »), а все bold-ран после него — модель
    arm_model_runs = []
    found_label = False
    for r in cell_10.paragraphs[5].runs:
        if not found_label and not r.bold and r.text.strip().endswith("՝"):
            found_label = True
            continue
        if found_label and r.bold:
            arm_model_runs.append(r)
    if arm_model_runs:
        arm_model_runs[0].text = model
        for r in arm_model_runs[1:]:
            r.text = ""

    # Серийный номер (армянская сторона) — paragraph[6]
    for run in cell_10.paragraphs[6].runs:
        if run.text == "PW0H2EGW":
            run.text = serial

    # ── Cell [1,1] — русская часть (ФИО + модель + серийник + пол) ──
    cell_11 = table.cell(1, 1)

    # Гражданин/гражданка РФ или РА
    for run in cell_11.paragraphs[1].runs:
        if "гражданин(ка) Р" in run.text:
            run.text = run.text.replace("гражданин(ка) Р", f"{citizen_word} ")
        # Ф → РФ или А → РА во втором run
        if run.text == "Ф":
            run.text = citizenship  # "РФ" или "РА"
        # Фамилия Имя Отчество
        if run.text == "Ситраков":
            run.text = rus_last
        elif run.text == " Леонид Станиславович":
            run.text = f" {rus_first} {rus_patr}"

    # Модель (русская сторона) — paragraph[5], bold-ран между «Модель: » и «Серийный номер: »
    rus_model_runs = []
    in_model_zone = False
    for r in cell_11.paragraphs[5].runs:
        if not in_model_zone and not r.bold and ("Модель" in r.text or ":" in r.text):
            continue
        if not in_model_zone and r.bold:
            in_model_zone = True
        if in_model_zone and r.bold:
            rus_model_runs.append(r)
        if in_model_zone and not r.bold and "Серийный" in r.text:
            break
    if rus_model_runs:
        rus_model_runs[0].text = model
        for r in rus_model_runs[1:]:
            r.text = ""

    # Серийный номер (русская сторона)
    for run in cell_11.paragraphs[5].runs:
        if run.text == "PW0H2EGW":
            run.text = serial

    doc.save(output_path)


# ─── GUI ────────────────────────────────────────────────────────────────────
class App:
    def __init__(self, root):
        import tkinter as tk
        from tkinter import ttk
        self.tk = tk
        self.ttk = ttk
        self.root = root
        root.title("Генерация акта приёма-передачи техники")
        root.geometry("680x660")
        root.resizable(False, False)

        # Принудительно UTF-8 для корректной работы IME/мнемонических раскладок
        root.tk.call('encoding', 'system', 'utf-8')

        # Стиль
        style = ttk.Style()
        style.theme_use("vista" if "vista" in style.theme_names() else "clam")

        # Загружаем модели
        self.models = load_models()

        # ── Переменные ──
        self.name_rus_var = StringVar()
        self.gender_var = StringVar(value="Мужской")
        self.citizenship_var = StringVar(value="РФ")
        self.model_var = StringVar()
        self.serial_var = StringVar()
        self.date_var = StringVar()

        # Заполняем дату
        self._update_date()

        # ── Интерфейс ── (используем ttk для современного вида)
        pad = {"padx": 15, "pady": 6}

        # Заголовок
        title = ttk.Label(root, text="📄 Генерация акта приёма-передачи",
                          font=("Segoe UI", 16, "bold"))
        title.pack(pady=(20, 10))

        main_frame = ttk.Frame(root)
        main_frame.pack(fill="both", expand=True, padx=20, pady=5)

        # ── Рамка: Данные сотрудника ──
        emp_frame = ttk.LabelFrame(main_frame, text="👤 Данные сотрудника", padding=12)
        emp_frame.pack(fill="x", pady=(0, 10))

        ttk.Label(emp_frame, text="ФИО (армянский):", font=("Segoe UI", 10)).grid(
            row=0, column=0, sticky="w", pady=3)
        # Text — собственный рендеринг Tk, отлично работает с Unicode
        self.name_arm_text = Text(
            emp_frame, height=1, wrap="none",
            font=("Sylfaen", 13), padx=4, pady=2,
            relief="solid", bd=1
        )
        self.name_arm_text.grid(row=0, column=1, sticky="ew", pady=3, padx=(5, 0))
        self.name_arm_text.bind("<Tab>", self._arm_text_tab)
        self.name_arm_text.bind("<Return>", self._arm_text_tab)

        ttk.Label(emp_frame, text="ФИО (русский):", font=("Segoe UI", 10)).grid(
            row=1, column=0, sticky="w", pady=3)
        Entry(emp_frame, textvariable=self.name_rus_var,
              font=("Segoe UI", 11), relief="solid", bd=1).grid(
                  row=1, column=1, sticky="ew", pady=3, padx=(5, 0))

        ttk.Label(emp_frame, text="Пол:", font=("Segoe UI", 10)).grid(
            row=2, column=0, sticky="w", pady=3)
        gender_frame = ttk.Frame(emp_frame)
        gender_frame.grid(row=2, column=1, sticky="w", pady=3, padx=(5, 0))
        ttk.Radiobutton(gender_frame, text="Мужской", variable=self.gender_var,
                        value="Мужской").pack(side="left")
        ttk.Radiobutton(gender_frame, text="Женский", variable=self.gender_var,
                        value="Женский").pack(side="left", padx=(15, 0))

        ttk.Label(emp_frame, text="Гражданство:", font=("Segoe UI", 10)).grid(
            row=3, column=0, sticky="w", pady=3)
        cit_frame = ttk.Frame(emp_frame)
        cit_frame.grid(row=3, column=1, sticky="w", pady=3, padx=(5, 0))
        ttk.Radiobutton(cit_frame, text="🇷🇺 РФ", variable=self.citizenship_var,
                        value="РФ").pack(side="left")
        ttk.Radiobutton(cit_frame, text="🇦🇲 РА", variable=self.citizenship_var,
                        value="РА").pack(side="left", padx=(15, 0))

        emp_frame.columnconfigure(1, weight=1)

        # ── Рамка: Данные ноутбука ──
        laptop_frame = ttk.LabelFrame(main_frame, text="💻 Данные ноутбука", padding=12)
        laptop_frame.pack(fill="x", pady=(0, 10))

        ttk.Label(laptop_frame, text="Модель:", font=("Segoe UI", 10)).grid(
            row=0, column=0, sticky="w", pady=3)

        model_row = ttk.Frame(laptop_frame)
        model_row.grid(row=0, column=1, sticky="ew", pady=3, padx=(5, 0))
        model_row.columnconfigure(0, weight=1)

        if self.models:
            self.model_var.set(self.models[0])
        self.model_menu = ttk.Combobox(model_row, textvariable=self.model_var,
                                       values=self.models, font=("Segoe UI", 10),
                                       state="normal")
        self.model_menu.pack(fill="x", expand=True, side="left")
        self.model_menu.bind("<<ComboboxSelected>>", self._on_model_selected)

        ttk.Button(model_row, text="➕", width=3,
                   command=self._add_new_model).pack(side="right", padx=(5, 0))

        ttk.Label(laptop_frame, text="Серийный номер:", font=("Segoe UI", 10)).grid(
            row=1, column=0, sticky="w", pady=3)
        ttk.Entry(laptop_frame, textvariable=self.serial_var,
                  font=("Consolas", 11)).grid(row=1, column=1, sticky="ew", pady=3, padx=(5, 0))

        ttk.Label(laptop_frame, text="Дата выдачи:", font=("Segoe UI", 10)).grid(
            row=2, column=0, sticky="w", pady=3)
        date_entry = ttk.Entry(laptop_frame, textvariable=self.date_var,
                               font=("Segoe UI", 11), state="readonly")
        date_entry.grid(row=2, column=1, sticky="ew", pady=3, padx=(5, 0))

        laptop_frame.columnconfigure(1, weight=1)

        # ── Кнопка генерации ──
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill="x", pady=(10, 5))

        self.generate_btn = ttk.Button(
            btn_frame, text="📄 Сгенерировать акт",
            command=self._generate, style="Accent.TButton"
        )
        self.generate_btn.pack(side="right")

        # Для акцентной кнопки
        style.configure("Accent.TButton", font=("Segoe UI", 11, "bold"))

        # Статус-бар
        self.status_var = StringVar(value="✅ Готов к работе")
        status_bar = ttk.Label(root, textvariable=self.status_var,
                               relief="sunken", anchor="w",
                               font=("Segoe UI", 9))
        status_bar.pack(fill="x", side="bottom", ipady=2)

        # ── Подсказка с форматами ──
        hint = ttk.Label(
            main_frame,
            text=("💡 Формат ввода:\n"
                  "  Армянский: Անուն Հայրանուն Ազգանուն  (пример: Լեոնիդ Ստանիսլավի Սիտրակով)\n"
                  "  Русский:   Фамилия Имя Отчество       (пример: Ситраков Леонид Станиславович)"),
            font=("Segoe UI", 9), foreground="#555",
            justify="left"
        )
        hint.pack(fill="x", pady=(10, 0))

    # ── Методы ──────────────────────────────────────────────────────────────

    def _get_name_arm(self):
        """Получить текст из Text-поля армянского ФИО."""
        return self.name_arm_text.get(1.0, "end-1c").strip()

    def _arm_text_tab(self, event):
        """Tab/Enter → переход к следующему полю."""
        event.widget.tk_focusNext().focus_set()
        return "break"

    def _update_date(self):
        """Обновляет поле с датой на сегодняшнюю."""
        today = date.today()
        months_ru = [
            "января", "февраля", "марта", "апреля", "мая", "июня",
            "июля", "августа", "сентября", "октября", "ноября", "декабря"
        ]
        self.date_var.set(f'{today.day} {months_ru[today.month - 1]} {today.year} г.')

    def _on_model_selected(self, event=None):
        """При выборе модели из списка."""
        pass

    def _add_new_model(self):
        """Добавить новую модель в список."""
        dialog = Tk()
        dialog.title("Новая модель")
        dialog.geometry("400x120")
        dialog.resizable(False, False)
        dialog.transient(self.root)
        dialog.grab_set()

        ttk.Label(dialog, text="Введите название модели:", font=("Segoe UI", 10)).pack(
            padx=15, pady=(15, 5), anchor="w")

        new_model_var = StringVar()
        entry = ttk.Entry(dialog, textvariable=new_model_var, font=("Segoe UI", 10))
        entry.pack(padx=15, fill="x", pady=(0, 10))
        entry.focus_set()

        def save():
            val = new_model_var.get().strip()
            if val:
                if val not in self.models:
                    self.models.append(val)
                    self.models.sort()
                    save_models(self.models)
                    self.model_menu["values"] = self.models
                self.model_var.set(val)
            dialog.destroy()

        def on_enter(event):
            save()

        entry.bind("<Return>", on_enter)

        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(fill="x", padx=15, pady=(0, 10))
        ttk.Button(btn_frame, text="Отмена", command=dialog.destroy).pack(side="left")
        ttk.Button(btn_frame, text="Сохранить", command=save).pack(side="right")

        dialog.wait_window()

    def _generate(self):
        """Генерирует Word-документ."""
        # ── Валидация ──
        name_arm = self._get_name_arm()
        name_rus = self.name_rus_var.get().strip()
        gender = self.gender_var.get()
        model = self.model_var.get().strip()
        serial = self.serial_var.get().strip()

        if not name_arm:
            messagebox.showerror("Ошибка", "Введите ФИО на армянском!")
            self.name_arm_text.focus_set()
            return
        if not name_rus:
            messagebox.showerror("Ошибка", "Введите ФИО на русском!")
            return
        if not model:
            messagebox.showerror("Ошибка", "Выберите или введите модель ноутбука!")
            return
        if not serial:
            messagebox.showerror("Ошибка", "Введите серийный номер!")
            return

        # Проверка количества слов в ФИО
        if len(name_arm.split()) < 3:
            messagebox.showerror("Ошибка",
                                 "Армянское ФИО должно содержать 3 части:\n"
                                 "Անուն Հայրանուն Ազգանուն")
            return
        if len(name_rus.split()) < 3:
            messagebox.showerror("Ошибка",
                                 "Русское ФИО должно содержать 3 части:\n"
                                 "Фамилия Имя Отчество")
            return

        # ── Создаём папку ──
        # Извлекаем фамилию для имени папки (из русского ФИО — первый элемент)
        rus_parts = name_rus.split()
        folder_name = f"{rus_parts[0]}_{date.today().strftime('%Y-%m-%d')}"
        output_subdir = os.path.join(OUTPUT_DIR, folder_name)
        os.makedirs(output_subdir, exist_ok=True)

        filename = f"Акт_приема_передачи_{rus_parts[0]}_{date.today().strftime('%Y-%m-%d')}.docx"
        output_path = os.path.join(output_subdir, filename)

        # ── Генерация ──
        try:
            generate_act(name_arm, name_rus, gender, self.citizenship_var.get(),
                         model, serial, output_path)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сгенерировать документ:\n{e}")
            return

        # Сохраняем модель, если её ещё нет
        if model not in self.models:
            self.models.append(model)
            self.models.sort()
            save_models(self.models)
            self.model_menu["values"] = self.models

        # Обновляем статус
        self.status_var.set(f"✅ Акт сохранён: {output_path}")
        messagebox.showinfo(
            "Готово",
            f"Документ сохранён:\n{output_path}"
        )


# ─── Запуск ─────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    root = Tk()
    app = App(root)
    root.mainloop()
